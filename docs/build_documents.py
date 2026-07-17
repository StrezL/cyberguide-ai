from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent
PURPLE = RGBColor(112, 66, 148)
DARK = RGBColor(35, 26, 44)
MUTED = RGBColor(104, 94, 112)
LIGHT = "F3ECF8"
PALE = "FAF7FC"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc: Document, running_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = DARK
    title.paragraph_format.space_after = Pt(7)
    title_ppr = title.element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    for name, size, before, after in (
        ("Heading 1", 17, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11.5, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = PURPLE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = running_title.upper()
    header.style = styles["Caption"]
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = MUTED
    header.runs[0].font.bold = True
    header.paragraph_format.space_after = Pt(0)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document, kind: str, title: str, subtitle: str) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run(kind.upper())
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = PURPLE
    kicker.paragraph_format.space_after = Pt(18)

    heading = doc.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(title)
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run = subtitle_p.add_run(subtitle)
    run.font.name = "Aptos"
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    accent = doc.add_paragraph()
    accent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    block = accent.add_run("  AI • SECURITY AWARENESS • CLOUD  ")
    block.font.name = "Aptos"
    block.font.size = Pt(9)
    block.font.bold = True
    block.font.color.rgb = WHITE
    block._element.get_or_add_rPr().append(_run_shading("7A4E9D"))
    accent.paragraph_format.space_after = Pt(70)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.line_spacing = 1.5
    r = meta.add_run("Prepared by: [Student Name]\nCourse: Gen AI & Cloud Computing\nDate: July 2026")
    r.font.name = "Aptos"
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    url = doc.add_paragraph()
    url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url.paragraph_format.space_before = Pt(30)
    r = url.add_run("LIVE AWS APPLICATION URL: [ADD AFTER DEPLOYMENT]")
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = PURPLE


def _run_shading(fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def add_lead(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = PURPLE
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc: Document, items) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_key_value_table(doc: Document, rows) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.75)
    header = table.rows[0]
    header.cells[0].text = "Item"
    header.cells[1].text = "Details"
    set_repeat_table_header(header)
    for cell in header.cells:
        set_cell_shading(cell, "7A4E9D")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.75)
        cells[1].width = Inches(4.75)
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], PALE)
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[0].paragraphs[0].runs[0].font.color.rgb = PURPLE
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.75 if cell is row.cells[0] else 4.75)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def concept_note() -> Path:
    doc = Document()
    configure_document(doc, "CyberGuide AI | Concept Note")
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.space_after = Pt(4)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.03
    for style_name in ("List Bullet", "List Number"):
        doc.styles[style_name].paragraph_format.space_after = Pt(3)
        doc.styles[style_name].paragraph_format.line_spacing = 1.05
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(11)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(5)
    add_cover(doc, "Project Concept Note", "CyberGuide AI", "Intelligent Security Awareness Platform")
    doc.add_page_break()

    doc.add_heading("1. Project Overview", level=1)
    add_lead(doc, "Concept", "A responsive AI-powered web application that helps learners understand digital threats, information protection, cloud security, and practical defensive measures through explanations, quizzes, and guided learning.")
    add_key_value_table(doc, [
        ("Application name", "CyberGuide AI"),
        ("Application type", "Intelligent Security Awareness Platform"),
        ("Primary audience", "Students, beginners, general internet users, and early-career professionals"),
        ("Backend", "Python FastAPI with streamed HTTP responses"),
        ("LLM API", "Groq API using the Llama 3.3 70B Versatile model"),
        ("Cloud deployment", "Docker container deployed through AWS App Runner"),
    ])

    doc.add_heading("2. Problem Statement", level=1)
    doc.add_paragraph("Digital users regularly encounter phishing, scams, weak authentication, insecure cloud practices, malware, and social-engineering attempts. Security guidance is often fragmented or too technical for beginners. This can prevent learners from understanding how threats work and which defensive actions are appropriate.")
    doc.add_paragraph("CyberGuide AI addresses this gap by providing a single, approachable platform where users can explore defined security domains and receive AI-generated explanations adapted to their level.")

    doc.add_heading("3. Objective", level=1)
    add_bullets(doc, [
        "Provide simple, accurate explanations of security concepts and internet threats.",
        "Connect each threat with practical prevention, detection, response, and recovery measures.",
        "Support active learning through AI-generated quizzes and examples.",
        "Deliver responses progressively through a streaming interface.",
        "Demonstrate secure API integration, containerization, and public AWS deployment.",
    ])

    doc.add_heading("4. Target Users and Use Cases", level=1)
    add_key_value_table(doc, [
        ("Students", "Revise core concepts, request simple explanations, and generate quick quizzes."),
        ("General users", "Learn how to recognize scams, phishing, malware indicators, and risky online behavior."),
        ("Beginners", "Select a learning level and receive accessible guidance without excessive jargon."),
        ("Early-career users", "Explore information, network, and cloud-security fundamentals using practical examples."),
    ])

    doc.add_heading("5. Key Application Features", level=1)
    add_bullets(doc, [
        "Seven learning domains: Internet Threats, Information Security, Cloud Security, Network Security, Social Engineering, Malware, and Defensive Measures.",
        "Three learning levels: beginner, intermediate, and advanced.",
        "Three output modes: concept explanation, defensive measures, and quiz generation.",
        "Suggested questions and domain cards that prepare focused prompts.",
        "Real-time streamed AI responses with copy functionality.",
        "Responsive lilac-themed interface optimized for desktop and mobile.",
        "Defensive-use guardrails that redirect unsafe requests toward safe learning.",
    ])

    doc.add_heading("6. Expected User Experience and Outcomes", level=1)
    doc.add_paragraph("A user opens the HTTPS application, selects a security domain and learning preference, enters a question, and receives an answer progressively. The interface clearly communicates that the application is intended for defensive education. By the end of a session, the learner should understand the selected threat or concept, recognize common warning signs, and identify realistic protective measures.")

    doc.add_heading("7. Success Criteria", level=1)
    add_bullets(doc, [
        "The application loads successfully on desktop and mobile browsers.",
        "The server keeps the LLM API key outside frontend code and version control.",
        "AI responses appear progressively and remain focused on safe defensive learning.",
        "The Docker image starts reliably and the health endpoint reports a healthy status.",
        "The final AWS URL is publicly accessible over HTTPS at submission time.",
    ])

    path = OUTPUT / "CyberGuide_AI_Concept_Note.docx"
    doc.save(path)
    return path


def project_report() -> Path:
    doc = Document()
    configure_document(doc, "CyberGuide AI | Project Report")
    add_cover(doc, "Final Project Report", "CyberGuide AI", "Vibe Coding: Building and Deploying an AI Web Application on AWS")
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    add_lead(doc, "Project outcome", "CyberGuide AI is a full-stack security-awareness application with a responsive frontend, FastAPI backend, streamed Groq LLM integration, Docker packaging, and an AWS-ready deployment design.")
    doc.add_paragraph("The project demonstrates an AI-first development workflow in which the learner acts as architect, designer, tester, and prompt director. AI assistance was used to explore the concept, refine the interface, draft code, identify errors, and improve documentation. Final decisions about scope, safety, architecture, and user experience were retained by the learner.")

    doc.add_heading("1. Application Overview", level=1)
    doc.add_paragraph("CyberGuide AI helps users learn about digital threats and protective measures through an interactive AI assistant. Instead of presenting only an open chat box, the application organizes learning into seven visible domains. Users can select their level and choose whether they want an explanation, defensive measures, or a quiz.")
    add_key_value_table(doc, [
        ("Frontend", "Semantic HTML, responsive CSS, and vanilla JavaScript"),
        ("Backend", "Python 3.12, FastAPI, Pydantic validation, and HTTPX"),
        ("AI model", "Llama 3.3 70B Versatile through the Groq API"),
        ("Streaming", "FastAPI StreamingResponse and browser ReadableStream"),
        ("Container", "Docker using a minimal Python base image and non-root user"),
        ("Cloud", "Amazon ECR image repository and AWS App Runner service"),
    ])

    doc.add_heading("2. Problem and Design Rationale", level=1)
    doc.add_paragraph("Many users know that online threats exist but struggle to distinguish phishing, malware, social engineering, network risks, and cloud-security responsibilities. The application therefore combines topic navigation with conversational AI. The lilac-and-orchid visual identity avoids common hacker-style clichés while retaining a serious, modern security character.")

    doc.add_heading("3. Prompting Strategy", level=1)
    doc.add_heading("3.1 System Prompt Framework", level=2)
    doc.add_paragraph("The backend supplies a fixed system prompt before each user request. It defines the AI role, educational scope, safety boundary, response style, and redirection behavior. The prompt requires defensive guidance and refuses instructions that could enable credential theft, malware, unauthorized access, evasion, exploitation, or destructive activity.")
    doc.add_heading("3.2 Structured User Context", level=2)
    doc.add_paragraph("Each request is transformed into a structured prompt containing four fields: learning domain, learner level, response mode, and user question. This makes the model response more consistent than sending the raw question alone.")
    doc.add_heading("3.3 Sample Development Prompts", level=2)
    add_bullets(doc, [
        "Design a responsive security-awareness platform with a refined lilac theme, seven learning domains, and an accessible AI question composer.",
        "Create a FastAPI endpoint that validates a security-learning request and streams an external LLM response without exposing the API key to the browser.",
        "Refine the system prompt so unsafe offensive requests are refused and redirected toward defensive concepts or safe lab learning.",
        "Package the FastAPI application in a small Docker image that listens on port 8080 and runs as a non-root user.",
        "Review the mobile layout for touch targets, readability, wrapping, keyboard interaction, and reduced-motion accessibility.",
    ])
    doc.add_heading("3.4 Example Runtime Prompt", level=2)
    add_lead(doc, "Example", "Learning domain: Cloud Security | Learner level: beginner | Response mode: measures | User question: How can a small business prevent accidental exposure of cloud data?")

    doc.add_heading("4. Phase-by-Phase Development", level=1)
    add_numbered(doc, [
        "Concept definition — selected a security-awareness platform aligned with academic and professional interests, then defined the target users and seven learning domains.",
        "Experience design — compared visual directions and selected a dark glass interface, then replaced blue accents with a lilac, violet, and orchid palette.",
        "Frontend development — built responsive navigation, topic cards, learning controls, question composer, streaming answer panel, copy action, and mobile breakpoints.",
        "Backend development — created validated FastAPI routes, the AI system prompt, structured prompt construction, Groq streaming, safe error handling, and a health endpoint.",
        "Containerization — created a Dockerfile, excluded secrets and unnecessary files, exposed port 8080, and configured a non-root runtime user.",
        "Testing — compiled the Python and JavaScript, ran automated endpoint tests, and verified the homepage, health check, validation, and unconfigured-key behavior.",
        "Cloud deployment — push the image to Amazon ECR, create an App Runner service, add the API key securely, verify HTTPS access, and insert the final URL into both documents.",
    ])

    doc.add_heading("5. Application Architecture", level=1)
    add_key_value_table(doc, [
        ("1. User interface", "Captures the domain, level, mode, and question; displays streamed text without storing the secret key."),
        ("2. FastAPI server", "Validates input, adds safety and learning instructions, and initiates the LLM request."),
        ("3. Groq LLM API", "Generates the response and returns incremental content chunks."),
        ("4. Streaming path", "The server forwards chunks to the browser, which appends them progressively to the answer panel."),
        ("5. AWS runtime", "App Runner pulls the Docker image from ECR, injects environment variables, and exposes the HTTPS service."),
    ])

    doc.add_heading("6. Security and Responsible-AI Measures", level=1)
    add_bullets(doc, [
        "The Groq API key is read only from the server environment and is excluded from Git and Docker build context.",
        "Pydantic restricts input length and accepts only defined learning levels and response modes.",
        "The system prompt establishes a defensive-only boundary and safe redirection policy.",
        "Upstream error details and credentials are not returned to the browser.",
        "The Docker container runs as an unprivileged user.",
        "The application does not require user accounts or collect personal information for its core function.",
    ])

    doc.add_heading("7. Challenges and Resolutions", level=1)
    add_key_value_table(doc, [
        ("Broad initial concept", "Converted a generic chatbot idea into defined learning domains and response modes."),
        ("Visual identity", "Replaced the initial blue cyber palette with a distinctive lilac system while maintaining contrast."),
        ("Streaming", "Used an asynchronous upstream stream and FastAPI StreamingResponse, then read chunks through the browser stream API."),
        ("Secret management", "Kept the key in environment variables and added both .gitignore and .dockerignore safeguards."),
        ("Unsafe questions", "Added a defensive system prompt that refuses harmful requests and redirects to safe concepts."),
        ("Deployment portability", "Used Docker and a health endpoint so the same application can run locally and on AWS."),
    ])

    doc.add_heading("8. Testing and Validation", level=1)
    doc.add_paragraph("Automated tests verified the homepage response, health-check JSON, and rejection of an empty question. Python compilation and JavaScript syntax checks completed successfully. A local service smoke test confirmed HTTP 200 for the homepage, a healthy status response, and an appropriate configuration message when the API key was absent.")
    add_bullets(doc, [
        "Test result: 3 automated tests passed.",
        "Health endpoint: GET /health returned {status: healthy}.",
        "Security check: no API key appears in HTML, CSS, JavaScript, Dockerfile, or committed configuration.",
        "Final deployment checks pending: live LLM response, public AWS HTTPS URL, and mobile-browser review.",
    ])

    doc.add_heading("9. Key Learnings and Reflection", level=1)
    doc.add_paragraph("This project demonstrated that vibe coding is not simply asking AI to generate an application. Useful results depended on clear requirements, iterative prompting, evaluation of design choices, validation of generated code, and understanding how each component communicates. The most important learning was connecting the visible user experience to backend security and cloud deployment decisions.")
    doc.add_paragraph("The project also reinforced that an AI feature needs boundaries. Because the application concerns security, defining a defensive educational purpose was as important as connecting the model. Containerization and the health endpoint showed how local application decisions affect the reliability of cloud deployment.")

    doc.add_heading("10. Final Deployment Checklist", level=1)
    add_bullets(doc, [
        "Create the Groq API key and test a live streamed response locally.",
        "Configure an AWS Budget alert before creating cloud resources.",
        "Build the Docker image and push it to Amazon ECR.",
        "Create the AWS App Runner service with port 8080 and /health monitoring.",
        "Add GROQ_API_KEY securely to the App Runner service configuration.",
        "Verify the application on desktop and mobile using its public HTTPS URL.",
        "Replace every [ADD AFTER DEPLOYMENT] marker with the final AWS URL.",
    ])

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("CyberGuide AI satisfies the planned full-stack design: a clear and responsive frontend, secure FastAPI backend, external LLM integration with real-time streaming, Docker containerization, and an AWS deployment architecture. Once the API key and AWS service are configured, the final public URL will complete the required deliverables.")

    path = OUTPUT / "CyberGuide_AI_Project_Report.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(concept_note())
    print(project_report())
